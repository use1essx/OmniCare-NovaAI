"""
File Parser for Questionnaire Generation
Parses various file types and extracts text content
Supports: PDF, Excel, CSV, Word, TXT, Images
"""

import io
import logging
from typing import Dict, Any
from pathlib import Path

# PDF parsing
try:
    import PyPDF2
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

# Excel parsing
try:
    import openpyxl  # noqa: F401
    import pandas as pd
    HAS_EXCEL = True
except ImportError:
    HAS_EXCEL = False

# Word parsing
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

logger = logging.getLogger(__name__)


class FileParser:
    """Parse different file types and extract text content"""
    
    ALLOWED_TYPES = [
        'pdf', 'xlsx', 'xls', 'csv', 'docx', 'txt',
        'jpg', 'jpeg', 'png', 'gif', 'bmp', 'webp'
    ]
    
    MAX_SIZE_MB = 10
    
    @classmethod
    def validate_file(cls, file_name: str, file_size: int) -> Dict[str, Any]:
        """
        Validate file type and size
        
        Args:
            file_name: Name of the file
            file_size: Size in bytes
            
        Returns:
            Dictionary with validation result
        """
        file_ext = Path(file_name).suffix.lower().lstrip('.')
        file_size_mb = file_size / (1024 * 1024)
        
        if file_ext not in cls.ALLOWED_TYPES:
            return {
                "valid": False,
                "error": f"File type .{file_ext} is not supported. Allowed types: {', '.join(cls.ALLOWED_TYPES)}"
            }
        
        if file_size_mb > cls.MAX_SIZE_MB:
            return {
                "valid": False,
                "error": f"File size ({file_size_mb:.2f}MB) exceeds maximum allowed size of {cls.MAX_SIZE_MB}MB"
            }
        
        return {"valid": True}
    
    @classmethod
    async def parse_file(cls, file_content: bytes, file_name: str, mime_type: str) -> Dict[str, Any]:
        """
        Parse file and extract text content
        
        Args:
            file_content: File content as bytes
            file_name: Original file name
            mime_type: MIME type of the file
            
        Returns:
            Dictionary with extracted text and metadata
        """
        file_ext = Path(file_name).suffix.lower().lstrip('.')
        
        try:
            if file_ext == 'pdf':
                return await cls._parse_pdf(file_content, file_name)
            elif file_ext in ['xlsx', 'xls']:
                return await cls._parse_excel(file_content, file_name)
            elif file_ext == 'csv':
                return await cls._parse_csv(file_content, file_name)
            elif file_ext == 'docx':
                return await cls._parse_docx(file_content, file_name)
            elif file_ext == 'txt':
                return await cls._parse_text(file_content, file_name)
            elif file_ext in ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'webp']:
                # Images handled by OCR (Gemini)
                return {
                    "text": "",
                    "metadata": {
                        "type": "image",
                        "note": "Image processing handled by AI OCR"
                    }
                }
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
                
        except Exception as e:
            logger.error(f"Error parsing file {file_name}: {e}")
            raise Exception(f"Error parsing file: {str(e)}")
    
    @classmethod
    async def _parse_pdf(cls, content: bytes, file_name: str) -> Dict[str, Any]:
        """Parse PDF file and extract text"""
        if not HAS_PDF:
            raise ImportError("PyPDF2 not installed. Run: pip install PyPDF2")
        
        try:
            logger.info(f"📄 Parsing PDF: {file_name}")
            
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            
            full_text = "\n\n".join(text_parts)
            
            logger.info(f"✅ PDF parsed - Pages: {len(pdf_reader.pages)}, Text length: {len(full_text)}")
            
            return {
                "text": full_text,
                "metadata": {
                    "pages": len(pdf_reader.pages),
                    "file_name": file_name
                }
            }
            
        except Exception as e:
            logger.error(f"Error parsing PDF: {e}")
            raise
    
    @classmethod
    async def _parse_excel(cls, content: bytes, file_name: str) -> Dict[str, Any]:
        """Parse Excel file and extract text"""
        if not HAS_EXCEL:
            raise ImportError("openpyxl and pandas not installed. Run: pip install openpyxl pandas")
        
        try:
            logger.info(f"📊 Parsing Excel: {file_name}")
            
            excel_file = io.BytesIO(content)
            
            # Try to read all sheets
            excel_data = pd.read_excel(excel_file, sheet_name=None)
            
            text_parts = []
            sheet_names = []
            
            for sheet_name, df in excel_data.items():
                sheet_names.append(sheet_name)
                text_parts.append(f"\n\nSheet: {sheet_name}\n")
                
                # Convert dataframe to text
                sheet_text = df.to_string(index=False)
                text_parts.append(sheet_text)
            
            full_text = "\n".join(text_parts)
            
            logger.info(f"✅ Excel parsed - Sheets: {len(sheet_names)}, Text length: {len(full_text)}")
            
            return {
                "text": full_text,
                "metadata": {
                    "sheets": len(sheet_names),
                    "sheet_names": sheet_names,
                    "file_name": file_name
                }
            }
            
        except Exception as e:
            logger.error(f"Error parsing Excel: {e}")
            raise
    
    @classmethod
    async def _parse_csv(cls, content: bytes, file_name: str) -> Dict[str, Any]:
        """Parse CSV file and extract text"""
        if not HAS_EXCEL:
            raise ImportError("pandas not installed. Run: pip install pandas")
        
        try:
            logger.info(f"📋 Parsing CSV: {file_name}")
            
            csv_file = io.BytesIO(content)
            df = pd.read_csv(csv_file)
            
            # Convert to text
            text = df.to_string(index=False)
            
            logger.info(f"✅ CSV parsed - Rows: {len(df)}, Text length: {len(text)}")
            
            return {
                "text": text,
                "metadata": {
                    "rows": len(df),
                    "columns": len(df.columns),
                    "column_names": list(df.columns),
                    "file_name": file_name
                }
            }
            
        except Exception as e:
            logger.error(f"Error parsing CSV: {e}")
            raise
    
    @classmethod
    async def _parse_docx(cls, content: bytes, file_name: str) -> Dict[str, Any]:
        """Parse Word document and extract text"""
        if not HAS_DOCX:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        try:
            logger.info(f"📝 Parsing Word document: {file_name}")
            
            docx_file = io.BytesIO(content)
            doc = Document(docx_file)
            
            # Extract text from paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            full_text = "\n\n".join(text_parts)
            
            logger.info(f"✅ Word document parsed - Paragraphs: {len(doc.paragraphs)}, Text length: {len(full_text)}")
            
            return {
                "text": full_text,
                "metadata": {
                    "paragraphs": len(doc.paragraphs),
                    "file_name": file_name
                }
            }
            
        except Exception as e:
            logger.error(f"Error parsing Word document: {e}")
            raise
    
    @classmethod
    async def _parse_text(cls, content: bytes, file_name: str) -> Dict[str, Any]:
        """Parse text file"""
        try:
            logger.info(f"📄 Parsing text file: {file_name}")
            
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            text = None
            
            for encoding in encodings:
                try:
                    text = content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                raise ValueError("Could not decode text file with any supported encoding")
            
            logger.info(f"✅ Text file parsed - Length: {len(text)}")
            
            return {
                "text": text,
                "metadata": {
                    "file_name": file_name
                }
            }
            
        except Exception as e:
            logger.error(f"Error parsing text file: {e}")
            raise


# Export parser instance
parser = FileParser()

