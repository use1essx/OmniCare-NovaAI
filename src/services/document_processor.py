"""
Document Processor Service

Handles document upload, text extraction, chunking, and embedding generation.
"""

import os
import hashlib
from typing import List, Dict, Any, Optional, Tuple, BinaryIO
from pathlib import Path
import mimetypes
from datetime import datetime

# Text extraction libraries
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

from src.core.logging import get_logger
from src.database.models_knowledge_base import KnowledgeDocument, KnowledgeChunk
from src.database.vector_store import get_vector_store
from src.services.embedding_service import get_embedding_service

logger = get_logger(__name__)


class DocumentProcessor:
    """Process documents for knowledge base"""
    
    def __init__(self):
        self.vector_store = get_vector_store()
        self.embedding_service = get_embedding_service()
        self.upload_dir = Path("./data/uploads")
        self.upload_dir.mkdir(parents=True, exist_ok=True)
    
    async def process_document(
        self,
        file: BinaryIO,
        filename: str,
        category: str,
        language: str = 'en',
        tags: Optional[List[str]] = None,
        uploaded_by: int = None,
        session: AsyncSession = None
    ) -> KnowledgeDocument:
        """
        Process uploaded document: extract text, chunk, embed, and store.
        
        Args:
            file: File object
            filename: Original filename
            category: Document category
            language: Document language
            tags: Optional tags
            uploaded_by: User ID who uploaded
            session: Database session
            
        Returns:
            KnowledgeDocument instance
        """
        logger.info(f"Processing document: {filename}")
        
        try:
            # 1. Save file
            file_path, file_size, file_type = await self._save_file(file, filename)
            
            # 2. Create database record
            doc = KnowledgeDocument(
                title=self._extract_title(filename),
                filename=filename,
                file_path=str(file_path),
                file_size=file_size,
                file_type=file_type,
                category=category,
                language=language,
                tags=tags or [],
                status='processing',
                uploaded_by=uploaded_by
            )
            
            if session:
                session.add(doc)
                await session.flush()  # Get doc.id
            
            # 3. Extract text
            text_content = await self._extract_text(file_path, file_type)
            
            if not text_content or len(text_content.strip()) < 50:
                doc.status = 'failed'
                doc.processing_error = 'Insufficient text content extracted'
                if session:
                    await session.commit()
                raise ValueError("Insufficient text content")
            
            # 4. Chunk text
            chunks = self._chunk_text(text_content)
            logger.info(f"Created {len(chunks)} chunks from document")
            
            # 5. Generate embeddings
            chunk_texts = [chunk['content'] for chunk in chunks]
            embeddings = await self.embedding_service.embed_batch(chunk_texts)
            
            # 6. Store in vector database
            chunk_ids = []
            metadatas = []
            
            for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                chunk_id = f"doc_{doc.id}_chunk_{i}"
                chunk_ids.append(chunk_id)
                
                metadata = {
                    'document_id': doc.id,
                    'document_title': doc.title,
                    'chunk_index': i,
                    'category': category,
                    'language': language,
                    'tags': tags or [],
                    'file_type': file_type,
                    **chunk.get('metadata', {})
                }
                metadatas.append(metadata)
            
            self.vector_store.add_documents(
                documents=chunk_texts,
                metadatas=metadatas,
                ids=chunk_ids,
                embeddings=embeddings
            )
            
            # 7. Create chunk records in database
            if session:
                for i, (chunk, chunk_id) in enumerate(zip(chunks, chunk_ids)):
                    db_chunk = KnowledgeChunk(
                        document_id=doc.id,
                        chunk_index=i,
                        content=chunk['content'],
                        content_hash=self._hash_content(chunk['content']),
                        chunk_metadata=chunk.get('metadata', {}),
                        token_count=chunk.get('token_count', 0),
                        vector_id=chunk_id,
                        embedding_model=self.embedding_service.model_name
                    )
                    session.add(db_chunk)
            
            # 8. Update document status
            doc.status = 'approved'  # Auto-approve for now
            doc.total_chunks = len(chunks)
            doc.total_characters = len(text_content)
            doc.metadata = {
                'chunks_created': len(chunks),
                'embedding_model': self.embedding_service.model_name,
                'processed_at': datetime.utcnow().isoformat()
            }
            
            if session:
                await session.commit()
            
            logger.info(f"Successfully processed document: {filename} (ID: {doc.id})")
            return doc
            
        except Exception as e:
            logger.error(f"Error processing document {filename}: {e}")
            if session and doc:
                doc.status = 'failed'
                doc.processing_error = str(e)
                await session.commit()
            raise
    
    async def _save_file(
        self,
        file: BinaryIO,
        filename: str
    ) -> Tuple[Path, int, str]:
        """
        Save uploaded file to disk.
        
        Returns:
            Tuple of (file_path, file_size, file_type)
        """
        # Generate unique filename
        timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
        safe_filename = self._sanitize_filename(filename)
        unique_filename = f"{timestamp}_{safe_filename}"
        file_path = self.upload_dir / unique_filename
        
        # Save file
        content = file.read()
        file_size = len(content)
        
        with open(file_path, 'wb') as f:
            f.write(content)
        
        # Detect file type
        file_type = mimetypes.guess_type(filename)[0] or 'application/octet-stream'
        
        return file_path, file_size, file_type
    
    async def _extract_text(self, file_path: Path, file_type: str) -> str:
        """Extract text from file based on type."""
        
        if 'pdf' in file_type.lower():
            return self._extract_pdf(file_path)
        elif 'word' in file_type.lower() or file_path.suffix == '.docx':
            return self._extract_docx(file_path)
        elif 'text' in file_type.lower() or file_path.suffix == '.txt':
            return self._extract_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    def _extract_pdf(self, file_path: Path) -> str:
        """Extract text from PDF."""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 not installed. Run: pip install PyPDF2")
        
        text = []
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page in pdf_reader.pages:
                text.append(page.extract_text())
        
        return '\n\n'.join(text)
    
    def _extract_docx(self, file_path: Path) -> str:
        """Extract text from DOCX."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        doc = DocxDocument(file_path)
        return '\n\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
    
    def _extract_txt(self, file_path: Path) -> str:
        """Extract text from TXT file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    def _chunk_text(
        self,
        text: str,
        chunk_size: int = 1000,
        chunk_overlap: int = 200
    ) -> List[Dict[str, Any]]:
        """
        Split text into overlapping chunks.
        
        Args:
            text: Text to chunk
            chunk_size: Target chunk size in characters
            chunk_overlap: Overlap between chunks
            
        Returns:
            List of chunk dicts with content and metadata
        """
        chunks = []
        
        # Simple character-based chunking
        # TODO: Implement smarter semantic chunking
        start = 0
        chunk_index = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence end in the last 100 chars
                search_start = max(start, end - 100)
                last_period = text.rfind('.', search_start, end)
                last_newline = text.rfind('\n', search_start, end)
                
                break_point = max(last_period, last_newline)
                if break_point > start:
                    end = break_point + 1
            
            chunk_text = text[start:end].strip()
            
            if chunk_text:
                chunks.append({
                    'content': chunk_text,
                    'metadata': {
                        'chunk_index': chunk_index,
                        'start_char': start,
                        'end_char': end
                    },
                    'token_count': len(chunk_text.split())  # Rough estimate
                })
                chunk_index += 1
            
            start = end - chunk_overlap
        
        return chunks
    
    def _extract_title(self, filename: str) -> str:
        """Extract title from filename."""
        # Remove extension and clean up
        title = Path(filename).stem
        title = title.replace('_', ' ').replace('-', ' ')
        return title.title()
    
    def _sanitize_filename(self, filename: str) -> str:
        """Sanitize filename for safe storage."""
        # Remove or replace unsafe characters
        safe_chars = set('abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789._-')
        return ''.join(c if c in safe_chars else '_' for c in filename)
    
    def _hash_content(self, content: str) -> str:
        """Generate hash of content for deduplication."""
        return hashlib.sha256(content.encode()).hexdigest()
    
    async def delete_document(
        self,
        document_id: int,
        session: AsyncSession
    ) -> bool:
        """
        Delete document and all associated chunks.
        
        Args:
            document_id: Document ID
            session: Database session
            
        Returns:
            True if successful
        """
        try:
            # Get document
            result = await session.execute(
                select(KnowledgeDocument).where(KnowledgeDocument.id == document_id)
            )
            doc = result.scalar_one_or_none()
            
            if not doc:
                return False
            
            # Delete from vector store
            self.vector_store.delete_by_metadata({'document_id': document_id})
            
            # Delete file
            if doc.file_path and os.path.exists(doc.file_path):
                os.remove(doc.file_path)
            
            # Delete from database (cascades to chunks)
            await session.delete(doc)
            await session.commit()
            
            logger.info(f"Deleted document ID: {document_id}")
            return True
            
        except Exception as e:
            logger.error(f"Error deleting document {document_id}: {e}")
            await session.rollback()
            return False


# Singleton instance
_document_processor_instance = None


def get_document_processor() -> DocumentProcessor:
    """Get or create document processor singleton."""
    global _document_processor_instance
    
    if _document_processor_instance is None:
        _document_processor_instance = DocumentProcessor()
    
    return _document_processor_instance
