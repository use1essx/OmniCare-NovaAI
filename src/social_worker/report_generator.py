"""
Report Generator for Social Worker Hub

Generates professional reports in PDF and DOCX formats:
- Progress reports
- Initial assessments
- Final reports
- Crisis reports
"""

import logging
from datetime import datetime
from pathlib import Path
from typing import Optional

from sqlalchemy import select

from .models import CaseFile, ProfessionalReport
from ..database.connection import get_async_db

logger = logging.getLogger(__name__)


# Try to import optional dependencies
try:
    from reportlab.lib import colors  # noqa: F401
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch  # noqa: F401
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle  # noqa: F401
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("reportlab not installed - PDF generation disabled")

try:
    from docx import Document
    from docx.shared import Inches, Pt  # noqa: F401
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed - DOCX generation disabled")


class ReportGenerator:
    """
    Generates professional reports for cases.
    
    Features:
    - PDF generation with reportlab
    - DOCX generation with python-docx
    - Multiple report types
    - Template support
    - Auto-numbering
    """
    
    REPORTS_DIR = Path("reports")
    
    def __init__(self, reports_dir: Optional[Path] = None):
        """Initialize report generator"""
        self.reports_dir = reports_dir or self.REPORTS_DIR
        self.reports_dir.mkdir(parents=True, exist_ok=True)
    
    async def generate_progress_report(
        self,
        case_id: int,
        generated_by: int,
        period_start: datetime,
        period_end: datetime,
        summary: str,
        progress_notes: Optional[str] = None,
        recommendations: Optional[str] = None,
        include_charts: bool = False,
        language: str = "en"
    ) -> ProfessionalReport:
        """
        Generate a progress report for a case.
        
        Args:
            case_id: Case ID
            generated_by: User ID generating report
            period_start: Report period start
            period_end: Report period end
            summary: Executive summary
            progress_notes: Progress notes
            recommendations: Recommendations
            include_charts: Include charts in report
            language: Report language
            
        Returns:
            ProfessionalReport record
        """
        async for session in get_async_db():
            try:
                # Get case details
                result = await session.execute(
                    select(CaseFile).where(CaseFile.id == case_id)
                )
                case = result.scalar_one_or_none()
                
                if not case:
                    raise ValueError(f"Case {case_id} not found")
                
                # Generate report number
                report_number = await self._generate_report_number(session)
                
                # Build report content
                content = self._build_progress_content(
                    case=case,
                    summary=summary,
                    progress_notes=progress_notes,
                    recommendations=recommendations,
                    period_start=period_start,
                    period_end=period_end,
                    language=language
                )
                
                # Generate title
                title = f"Progress Report - {case.case_number}"
                if language == "zh-HK":
                    title = f"進度報告 - {case.case_number}"
                
                # Generate PDF and DOCX
                pdf_path = None
                docx_path = None
                
                if PDF_AVAILABLE:
                    pdf_path = await self._generate_pdf(
                        report_number, title, content, case
                    )
                
                if DOCX_AVAILABLE:
                    docx_path = await self._generate_docx(
                        report_number, title, content, case
                    )
                
                # Create report record
                report = ProfessionalReport(
                    case_id=case_id,
                    child_id=case.child_id,
                    report_number=report_number,
                    report_type='progress',
                    title=title,
                    summary=summary,
                    content=content,
                    pdf_path=str(pdf_path) if pdf_path else None,
                    docx_path=str(docx_path) if docx_path else None,
                    includes_charts=include_charts,
                    includes_recommendations=recommendations is not None,
                    period_start=period_start.date(),
                    period_end=period_end.date(),
                    language=language,
                    generated_by=generated_by,
                    generation_method='manual',
                    review_status='draft'
                )
                
                session.add(report)
                await session.commit()
                await session.refresh(report)
                
                logger.info(f"Generated progress report {report_number}")
                return report
                
            except Exception as e:
                await session.rollback()
                logger.error(f"Error generating report: {e}")
                raise
    
    async def generate_initial_assessment(
        self,
        case_id: int,
        generated_by: int,
        assessment_summary: str,
        presenting_problems: str,
        background: str,
        risk_assessment: str,
        recommendations: str,
        language: str = "en"
    ) -> ProfessionalReport:
        """Generate initial assessment report"""
        async for session in get_async_db():
            try:
                result = await session.execute(
                    select(CaseFile).where(CaseFile.id == case_id)
                )
                case = result.scalar_one_or_none()
                
                if not case:
                    raise ValueError(f"Case {case_id} not found")
                
                report_number = await self._generate_report_number(session)
                
                # Build content
                content = f"""
# Initial Assessment Report

## Case Information
- Case Number: {case.case_number}
- Date: {datetime.utcnow().strftime('%Y-%m-%d')}

## Summary
{assessment_summary}

## Presenting Problems
{presenting_problems}

## Background
{background}

## Risk Assessment
{risk_assessment}

## Recommendations
{recommendations}
"""
                
                title = f"Initial Assessment - {case.case_number}"
                
                pdf_path = None
                if PDF_AVAILABLE:
                    pdf_path = await self._generate_pdf(
                        report_number, title, content, case
                    )
                
                report = ProfessionalReport(
                    case_id=case_id,
                    child_id=case.child_id,
                    report_number=report_number,
                    report_type='initial',
                    title=title,
                    summary=assessment_summary,
                    content=content,
                    pdf_path=str(pdf_path) if pdf_path else None,
                    includes_recommendations=True,
                    language=language,
                    generated_by=generated_by,
                    review_status='draft'
                )
                
                session.add(report)
                await session.commit()
                await session.refresh(report)
                
                return report
                
            except Exception as e:
                await session.rollback()
                logger.error(f"Error generating initial assessment: {e}")
                raise
    
    def _build_progress_content(
        self,
        case: CaseFile,
        summary: str,
        progress_notes: Optional[str],
        recommendations: Optional[str],
        period_start: datetime,
        period_end: datetime,
        language: str
    ) -> str:
        """Build progress report content in Markdown"""
        
        if language == "zh-HK":
            content = f"""
# 進度報告

## 個案資料
- 個案編號: {case.case_number}
- 報告期間: {period_start.strftime('%Y-%m-%d')} 至 {period_end.strftime('%Y-%m-%d')}
- 風險程度: {case.risk_level or 'N/A'}/100 ({case.risk_category or 'N/A'})
- 個案狀態: {case.status}

## 摘要
{summary}

## 進度記錄
{progress_notes or '未有記錄'}

## 建議
{recommendations or '未有建議'}

---
報告生成日期: {datetime.utcnow().strftime('%Y-%m-%d %H:%M')}
"""
        else:
            content = f"""
# Progress Report

## Case Information
- Case Number: {case.case_number}
- Period: {period_start.strftime('%Y-%m-%d')} to {period_end.strftime('%Y-%m-%d')}
- Risk Level: {case.risk_level or 'N/A'}/100 ({case.risk_category or 'N/A'})
- Status: {case.status}

## Executive Summary
{summary}

## Progress Notes
{progress_notes or 'No notes recorded'}

## Recommendations
{recommendations or 'No recommendations'}

---
Report Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M')}
"""
        
        return content
    
    async def _generate_pdf(
        self,
        report_number: str,
        title: str,
        content: str,
        case: CaseFile
    ) -> Optional[Path]:
        """Generate PDF report"""
        if not PDF_AVAILABLE:
            return None
        
        try:
            filename = f"{report_number}.pdf"
            filepath = self.reports_dir / filename
            
            doc = SimpleDocTemplate(
                str(filepath),
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )
            
            styles = getSampleStyleSheet()
            story = []
            
            # Title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=30
            )
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 12))
            
            # Content (basic markdown to paragraphs)
            for line in content.split('\n'):
                line = line.strip()
                if not line:
                    story.append(Spacer(1, 6))
                elif line.startswith('# '):
                    story.append(Paragraph(line[2:], styles['Heading1']))
                elif line.startswith('## '):
                    story.append(Paragraph(line[3:], styles['Heading2']))
                elif line.startswith('- '):
                    story.append(Paragraph(f"• {line[2:]}", styles['Normal']))
                else:
                    story.append(Paragraph(line, styles['Normal']))
            
            doc.build(story)
            
            logger.info(f"Generated PDF: {filepath}")
            return filepath
            
        except Exception as e:
            logger.error(f"Error generating PDF: {e}")
            return None
    
    async def _generate_docx(
        self,
        report_number: str,
        title: str,
        content: str,
        case: CaseFile
    ) -> Optional[Path]:
        """Generate DOCX report"""
        if not DOCX_AVAILABLE:
            return None
        
        try:
            filename = f"{report_number}.docx"
            filepath = self.reports_dir / filename
            
            doc = Document()
            
            # Title
            heading = doc.add_heading(title, 0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Content
            for line in content.split('\n'):
                line = line.strip()
                if not line:
                    doc.add_paragraph()
                elif line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('- '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                else:
                    doc.add_paragraph(line)
            
            doc.save(str(filepath))
            
            logger.info(f"Generated DOCX: {filepath}")
            return filepath
            
        except Exception as e:
            logger.error(f"Error generating DOCX: {e}")
            return None
    
    async def _generate_report_number(self, session) -> str:
        """Generate unique report number"""
        from sqlalchemy import func
        
        year = datetime.utcnow().year
        prefix = f"RPT-{year}-"
        
        result = await session.execute(
            select(func.max(ProfessionalReport.report_number)).where(
                ProfessionalReport.report_number.like(f"{prefix}%")
            )
        )
        max_number = result.scalar()
        
        if max_number:
            seq = int(max_number.split('-')[-1]) + 1
        else:
            seq = 1
        
        return f"{prefix}{seq:04d}"


# Singleton
_report_generator: Optional[ReportGenerator] = None


def get_report_generator() -> ReportGenerator:
    """Get or create report generator singleton"""
    global _report_generator
    if _report_generator is None:
        _report_generator = ReportGenerator()
    return _report_generator

